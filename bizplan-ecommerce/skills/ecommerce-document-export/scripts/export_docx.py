#!/usr/bin/env python3
"""
Export a markdown business plan to a professionally styled Word document.

Usage:
    python export_docx.py <input.md> <output.docx> --title "Business Name" [--date "March 2026"]

Preprocessing pipeline:
    1. Strip existing document header (replaced by YAML front matter)
    2. Promote headings (## N. Section → # Section for page breaks)
    3. Convert [N] citations to superscript ^[N]^
    3b. Convert [https://...] to [url](url) for clickable hyperlinks
    4. Ensure blank lines before lists (Pandoc requirement)
    5. Remove horizontal rules (--- separators)
    6. Generate YAML front matter with cover page and TOC

Post-processing:
    7. Style cover page (center, spacing) via python-docx
    8. Add visible borders to all tables via python-docx

Requires: pandoc, python-docx>=1.1.0
"""

import argparse
import re
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path


# ============================================================================
# Preprocessing
# ============================================================================


def strip_document_header(md: str) -> str:
    """
    Remove the existing markdown header block that gets replaced by YAML front matter.

    Strips patterns like:
        # Terra & Tide
        ## Business Plan — 2026-2031

        *Prepared March 2026*

        ---
    """
    # Match: H1 title, optional H2 subtitle, optional italicized date, optional HR
    md = re.sub(
        r'^#\s+.+\n'              # H1 title line
        r'(?:##\s+.+\n)?'         # Optional H2 subtitle line
        r'(?:\n\*[^*]+\*\n)?'     # Optional *italicized date*
        r'(?:\n---\n)?',          # Optional horizontal rule
        '',
        md,
    )
    return md.lstrip('\n')


def promote_headings(md: str) -> str:
    """
    Promote heading levels so main sections become H1 (triggering page breaks).

    Transforms:
        ## 1. Executive Summary  →  # Executive Summary
        ### What We Sell         →  ## What We Sell
        #### Detail              →  ### Detail

    Also strips section numbers (1., 2., etc.) from main section headings.
    """
    lines = []
    for line in md.split('\n'):
        # ## N. Section Name → # Section Name
        m = re.match(r'^##\s+\d+\.\s+(.+)$', line)
        if m:
            lines.append(f'# {m.group(1)}')
            continue

        # ### Subsection → ## Subsection
        m = re.match(r'^###\s+(.+)$', line)
        if m:
            lines.append(f'## {m.group(1)}')
            continue

        # #### Detail → ### Detail
        m = re.match(r'^####\s+(.+)$', line)
        if m:
            lines.append(f'### {m.group(1)}')
            continue

        lines.append(line)

    return '\n'.join(lines)


def convert_citations_to_superscript(md: str) -> str:
    """
    Convert [N] bracket citations to Pandoc superscript syntax ^[N]^.

    Only matches 1-3 digit numbers. Avoids markdown links [text](url)
    and 4-digit years like [2032].
    """
    return re.sub(r'\[(\d{1,3})\](?!\()', r'^\\[\1\\]^', md)


def convert_urls_to_links(md: str) -> str:
    """Convert [https://...] to [url](url) for Pandoc hyperlink generation."""
    return re.sub(r'\[(https?://[^\]]+)\](?!\()', r'[\1](\1)', md)


def ensure_blank_lines_before_lists(md: str) -> str:
    """Ensure blank lines before bullet/numbered lists (Pandoc requirement)."""
    lines = md.split('\n')
    result = []
    for i, line in enumerate(lines):
        is_list = bool(re.match(r'^[-*+] |^\d+\. ', line))
        if is_list and i > 0:
            prev = lines[i - 1]
            if prev.strip() != '' and not re.match(r'^[-*+] |^\d+\. ', prev):
                result.append('')
        result.append(line)
    return '\n'.join(result)


def remove_horizontal_rules(md: str) -> str:
    """Remove --- horizontal rules (not needed with H1 page breaks)."""
    # Avoid matching YAML front matter delimiters (handled separately)
    md = re.sub(r'\n\n---\n', '\n\n', md)
    md = re.sub(r'\n---\s*$', '', md)
    return md


def generate_front_matter(md: str, title: str, date: str) -> str:
    """Generate YAML front matter (cover page) and static Table of Contents."""
    lines = [
        "---",
        f'title: "{title}"',
        'subtitle: "Business Plan"',
        f'date: "{date}"',
        "---",
        "",
    ]

    # Extract H1 headings for TOC
    headings = re.findall(r'^#\s+(.+)$', md, re.MULTILINE)
    if headings:
        lines.extend(["# Table of Contents", ""])
        seen = set()
        for h in headings:
            clean = h.strip()
            if clean not in seen and clean != "Table of Contents":
                seen.add(clean)
                lines.append(f"- {clean}")
        lines.extend(["", ""])

    return '\n'.join(lines)


# ============================================================================
# Post-processing
# ============================================================================


def style_cover_page(docx_path: Path) -> None:
    """Post-process the docx to center and space the cover page elements."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document(docx_path)

    title_para = None
    subtitle_para = None
    date_para = None

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name == "Title" and title_para is None:
            title_para = para
        elif style_name == "Subtitle" and subtitle_para is None:
            subtitle_para = para
        elif style_name == "Date" and date_para is None:
            date_para = para

    if title_para:
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(180)

    if subtitle_para:
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.paragraph_format.space_before = Pt(12)

    if date_para:
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.paragraph_format.space_before = Pt(300)

    doc.save(docx_path)


def style_tables(docx_path: Path) -> None:
    """Post-process tables to add visible cell borders."""
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn

    doc = Document(docx_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                existing = tcPr.find(qn('w:tcBorders'))
                if existing is not None:
                    tcPr.remove(existing)
                tcBorders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    '</w:tcBorders>'
                )
                tcPr.append(tcBorders)
    doc.save(docx_path)


# ============================================================================
# Core Conversion
# ============================================================================


def count_words(text: str) -> int:
    """Count words excluding markdown syntax."""
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'[#*_`~]', '', text)
    return len(text.split())


def convert_markdown_to_docx(
    md_content: str,
    output_path: Path,
    document_title: str,
    date: str,
    reference_doc: Path,
) -> dict:
    """
    Full conversion pipeline: preprocess → Pandoc → post-process.

    Returns dict with output_path, word_count, page_estimate.
    """
    word_count = count_words(md_content)
    page_estimate = max(1, (word_count + 499) // 500)

    # 1. Strip existing document header
    md_content = strip_document_header(md_content)

    # 2. Promote headings (## N. → # for page breaks)
    md_content = promote_headings(md_content)

    # 3. Convert citations
    md_content = convert_citations_to_superscript(md_content)

    # 3b. Convert bracketed URLs to clickable links
    md_content = convert_urls_to_links(md_content)

    # 4. Ensure blank lines before lists
    md_content = ensure_blank_lines_before_lists(md_content)

    # 5. Remove horizontal rules
    md_content = remove_horizontal_rules(md_content)

    # 6. Generate front matter and TOC (must be after heading promotion)
    front_matter = generate_front_matter(md_content, document_title, date)
    final_md = front_matter + md_content

    # Write to temp file for Pandoc
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        mode='w', suffix='.md', delete=False, encoding='utf-8'
    ) as f:
        f.write(final_md)
        temp_path = Path(f.name)

    try:
        cmd = [
            "pandoc", "-f", "markdown", "-t", "docx",
            "-o", str(output_path),
        ]
        if reference_doc and reference_doc.exists():
            cmd.extend(["--reference-doc", str(reference_doc)])
        cmd.append(str(temp_path))

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            print(f"Pandoc error: {result.stderr}", file=sys.stderr)
            sys.exit(1)

        # 7. Post-process cover page and tables
        style_cover_page(output_path)
        style_tables(output_path)

        return {
            "output_path": str(output_path),
            "word_count": word_count,
            "page_estimate": page_estimate,
            "file_size_kb": round(output_path.stat().st_size / 1024, 1),
        }

    finally:
        if temp_path.exists():
            temp_path.unlink()


# ============================================================================
# CLI
# ============================================================================


def main():
    parser = argparse.ArgumentParser(description="Export markdown business plan to styled DOCX")
    parser.add_argument("input", help="Path to markdown business plan")
    parser.add_argument("output", help="Path for output .docx file")
    parser.add_argument("--title", required=True, help="Business name for cover page")
    parser.add_argument("--date", default=None, help="Date for cover page (default: current month)")
    parser.add_argument(
        "--reference-doc", default=None,
        help="Path to reference.docx (default: templates/reference.docx)",
    )
    args = parser.parse_args()

    # Resolve paths
    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    # Default reference doc
    if args.reference_doc:
        ref_doc = Path(args.reference_doc)
    else:
        ref_doc = Path(__file__).parent.parent / "templates" / "reference.docx"

    # Default date
    date = args.date or datetime.now().strftime("%B %Y")

    # Read input
    md_content = input_path.read_text(encoding="utf-8")

    # Check Pandoc
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, timeout=10, check=True)
    except (FileNotFoundError, subprocess.CalledProcessError):
        print("Error: Pandoc is not installed. Install with: brew install pandoc", file=sys.stderr)
        sys.exit(1)

    # Convert
    print(f"Exporting: {input_path} -> {output_path}")
    print(f"Title: {args.title}")
    print(f"Date: {date}")
    print(f"Reference doc: {ref_doc}")

    result = convert_markdown_to_docx(
        md_content=md_content,
        output_path=output_path,
        document_title=args.title,
        date=date,
        reference_doc=ref_doc,
    )

    print(f"\nExport complete:")
    print(f"  Output: {result['output_path']}")
    print(f"  Word count: {result['word_count']:,}")
    print(f"  Page estimate: ~{result['page_estimate']} pages")
    print(f"  File size: {result['file_size_kb']} KB")


if __name__ == "__main__":
    main()
