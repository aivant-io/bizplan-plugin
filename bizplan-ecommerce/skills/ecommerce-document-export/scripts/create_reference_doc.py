#!/usr/bin/env python3
"""
Generate the reference document for Pandoc styling.

Run this script to regenerate the reference.docx template:
    python create_reference_doc.py

The reference document defines heading styles, body text, cover page styles,
table styles, list formatting, and page layout used by Pandoc during conversion.
"""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


def _override_theme_font(style, font_name: str):
    """Force a style's font, overriding Word theme fonts via XML."""
    rPr = style._element.get_or_add_rPr()
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>'
    )
    rPr.append(rFonts)


def _add_hyperlink(paragraph, url: str, text: str):
    """Add a hyperlink to a paragraph (python-docx lacks native support)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def create_reference_document(output_path: Path) -> None:
    """Create a professionally styled reference document for Pandoc."""
    doc = Document()

    # Page Setup
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Color Palette
    PRIMARY = RGBColor(0x1A, 0x36, 0x5D)    # Dark navy
    SECONDARY = RGBColor(0x2E, 0x5C, 0x8A)  # Medium blue
    ACCENT = RGBColor(0x3D, 0x85, 0xC6)     # Light blue
    TEXT = RGBColor(0x33, 0x33, 0x33)        # Dark gray
    MUTED = RGBColor(0x66, 0x66, 0x66)       # Medium gray

    # Heading 1 — Main sections, page break before
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = PRIMARY
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True
    _override_theme_font(h1, "Times New Roman")

    # Heading 2 — Subsections
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = SECONDARY
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True
    _override_theme_font(h2, "Times New Roman")

    # Heading 3 — Minor headings
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = ACCENT
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True
    _override_theme_font(h3, "Times New Roman")

    # Heading 4 — Detail headings
    h4 = doc.styles["Heading 4"]
    h4.font.name = "Times New Roman"
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = TEXT
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)
    h4.paragraph_format.keep_with_next = True
    _override_theme_font(h4, "Times New Roman")

    # Normal — Body text
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    # Title — Cover page business name
    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(44)
    title.font.bold = True
    title.font.color.rgb = PRIMARY
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    _override_theme_font(title, "Times New Roman")

    # Subtitle — Cover page "Business Plan"
    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Times New Roman"
    subtitle.font.size = Pt(24)
    subtitle.font.bold = False
    subtitle.font.italic = True
    subtitle.font.color.rgb = SECONDARY
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    _override_theme_font(subtitle, "Times New Roman")

    # Date — Cover page date
    try:
        date_style = doc.styles.add_style("Date", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        date_style = doc.styles["Date"]
    date_style.font.name = "Times New Roman"
    date_style.font.size = Pt(14)
    date_style.font.color.rgb = TEXT
    date_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _override_theme_font(date_style, "Times New Roman")

    # TOC Heading
    if "TOC Heading" in doc.styles:
        toc = doc.styles["TOC Heading"]
        toc.font.name = "Times New Roman"
        toc.font.size = Pt(16)
        toc.font.bold = True
        toc.font.color.rgb = PRIMARY
        _override_theme_font(toc, "Times New Roman")

    # Hyperlink
    try:
        hl = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    except ValueError:
        hl = doc.styles["Hyperlink"]
    hl.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    hl.font.underline = True

    # Block Quote
    try:
        bq = doc.styles.add_style("Block Quote", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        bq = doc.styles["Block Quote"]
    bq.font.name = "Times New Roman"
    bq.font.size = Pt(11)
    bq.font.italic = True
    bq.font.color.rgb = MUTED
    bq.paragraph_format.left_indent = Inches(0.5)
    bq.paragraph_format.space_before = Pt(6)
    bq.paragraph_format.space_after = Pt(6)
    _override_theme_font(bq, "Times New Roman")

    # ---- Sample content so Pandoc learns all styles ----
    doc.add_paragraph("Business Plan", style="Title")
    doc.add_paragraph("Table of Contents", style="TOC Heading")
    doc.add_paragraph("(Generated automatically)", style="Normal")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This is the executive summary section. It provides a high-level overview "
        "of the business plan, key metrics, and strategic direction.",
        style="Normal",
    )

    doc.add_heading("Business Overview", level=1)
    doc.add_heading("Company Description", level=2)
    doc.add_paragraph("Description of the company and value proposition.", style="Normal")

    doc.add_heading("Products and Services", level=2)
    doc.add_paragraph("Overview of products and services offered.", style="Normal")

    doc.add_heading("Key Features", level=3)
    doc.add_paragraph("Detailed feature descriptions.", style="Normal")

    doc.add_heading("Technical Specifications", level=4)
    doc.add_paragraph("Technical details and specifications.", style="Normal")

    # List examples
    doc.add_paragraph("Key highlights:", style="Normal")
    for item in ["First point", "Second point", "Third point"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Strategic priorities:", style="Normal")
    for item in ["Priority one", "Priority two", "Priority three"]:
        doc.add_paragraph(item, style="List Number")

    # Hyperlink example
    doc.add_heading("References", level=2)
    p = doc.add_paragraph("For more information, see ", style="Normal")
    _add_hyperlink(p, "https://example.com", "Example Source")
    p.add_run(".")

    # Table example
    doc.add_heading("Financial Summary", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    headers[0].text = "Year"
    headers[1].text = "Revenue"
    headers[2].text = "Growth"
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
    for i, (y, r, g) in enumerate(
        [("2026", "$325,000", "-"), ("2027", "$650,000", "100%"), ("2028", "$1,300,000", "100%")],
        1,
    ):
        row = table.rows[i].cells
        row[0].text = y
        row[1].text = r
        row[2].text = g

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Created reference document: {output_path}")


if __name__ == "__main__":
    output = Path(__file__).parent.parent / "templates" / "reference.docx"
    create_reference_document(output)
